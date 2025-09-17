# app.py
import base64
import io
import os
import re
import sqlite3
import textwrap
import hashlib
from datetime import datetime, date, timedelta
from calendar import monthrange
from typing import List, Tuple, Optional, Dict, Any

from fastapi import FastAPI, Request, Form, UploadFile, File, Depends, HTTPException
from fastapi.responses import RedirectResponse, StreamingResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from starlette.middleware.sessions import SessionMiddleware
from starlette.middleware.base import BaseHTTPMiddleware

import requests
from PIL import Image, ImageOps
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------------------------------------------------------
# ENV & CONSTANTS
# -----------------------------------------------------------------------------
APP_NAME = os.getenv("APP_NAME", "Otel Planlama Stüdyosu")
SECRET_KEY = os.getenv("SECRET_KEY", "change-me-in-env")
USAGE_DB_PATH = os.getenv("USAGE_DB_PATH", "/tmp/usage.db")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", os.getenv("GeminiAPI", ""))
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.0-flash")  # 2.5 varsa kullanın
GEMINI_TIMEOUT = 40

ADMIN_CODE_USER = os.getenv("ADMIN_CODE_USER", "admin")
ADMIN_CODE_PASS = os.getenv("ADMIN_CODE_PASS", "admin123")

# DOCX yerleşim sabitleri
PAGE_MARGIN_CM = 1.27  # üst/alt/sol/sağ
IMG_W_CM = 16.0
IMG_H_CM = 20.0
FONT_BASE_PT = 10.0  # açıklama & hashtag
FONT_CONTACT_PT = 9.5
FONT_DATE_PT = 11.0

# -----------------------------------------------------------------------------
# APP
# -----------------------------------------------------------------------------
app = FastAPI(title=APP_NAME)
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)

# Statics & Templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


# -----------------------------------------------------------------------------
# DB HELPERS
# -----------------------------------------------------------------------------
def db() -> sqlite3.Connection:
    os.makedirs(os.path.dirname(USAGE_DB_PATH), exist_ok=True)
    conn = sqlite3.connect(USAGE_DB_PATH, check_same_thread=False)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_sha TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'user',
            created_at TEXT NOT NULL
        )
    """)
    conn.commit()
    return conn


def sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def seed_admin():
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT id FROM users WHERE username = ?", (ADMIN_CODE_USER,))
    if not cur.fetchone():
        cur.execute(
            "INSERT INTO users(username, password_sha, role, created_at) VALUES(?,?,?,?)",
            (ADMIN_CODE_USER, sha256(ADMIN_CODE_PASS), "admin", datetime.utcnow().isoformat()),
        )
        conn.commit()
    conn.close()


seed_admin()


# -----------------------------------------------------------------------------
# AUTH HELPERS
# -----------------------------------------------------------------------------
def current_user(request: Request) -> Optional[Dict[str, Any]]:
    return request.session.get("user")


def require_login(request: Request):
    user = current_user(request)
    if not user:
        raise HTTPException(status_code=307, detail="Redirect", headers={"Location": "/login"})
    return user


def require_admin(request: Request):
    user = require_login(request)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin required")
    return user


# -----------------------------------------------------------------------------
# MIDDLEWARE: Pretty errors for redirect exceptions
# -----------------------------------------------------------------------------
class RedirectMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        try:
            return await call_next(request)
        except HTTPException as exc:
            # capture our manual redirect from require_login
            if exc.status_code == 307 and exc.headers and "Location" in exc.headers:
                return RedirectResponse(url=exc.headers["Location"], status_code=307)
            raise

app.add_middleware(RedirectMiddleware)


# -----------------------------------------------------------------------------
# IMAGE & GEMINI
# -----------------------------------------------------------------------------
def open_image_safely(raw: bytes) -> Image.Image:
    img = Image.open(io.BytesIO(raw))
    img = ImageOps.exif_transpose(img).convert("RGB")
    return img


def gemini_json(image_bytes: bytes, hotel_info: str) -> Tuple[str, List[str]]:
    """
    Görsele uygun Türkçe açıklama (emoji'li) ve 4 hashtag döndürür.
    Gemini başarısız olursa basit bir yedek üretim yapar.
    """
    # İstek metni: sadece JSON döndürmesi için net talimat
    sys_prompt = (
        "Aşağıdaki görüntüye uygun, Instagram gönderisi için TÜRKÇE tek cümle "
        "bir açıklama üret. Emoji kullan (en az 2-3 emoji, ama abartma). "
        "Kaçamak, gibi kelimeler kullanma. "
        "Otelin ismini açıklamaya ekleme. "
        "3. çoğul şahıs ile yaz. "
        "Siz, biz gibi kurumsal dili koru "
        "Ayrıca görüntüye uygun tam 4 hashtag üret ve hepsi # ile başlasın. "
        "Cevabı JSON ver:\n"
        "{\n"
        '  "caption": "<emoji içeren tek cümle>",\n'
        '  "hashtags": ["#...", "#...", "#...", "#..."]\n'
        "}"
        "\nBaşlık etiketi, açıklama etiketi, iletişim vb. ekleme. Sadece JSON."
    )

    # Gemini yoksa fallback
    if not GEMINI_API_KEY:
        img = open_image_safely(image_bytes)
        w, h = img.size
        ratio = w / max(1, h)
        # basit renk tespiti
        small = img.resize((64, 64))
        avg = tuple(sum(p[i] for p in small.getdata()) // (64 * 64) for i in range(3))
        mood = "sıcak" if avg[0] > avg[2] else "serin"
        caption = f"{mood} tonların öne çıktığı bu karede keyif ve konfor bir arada 🌿✨"
        hashtags = ["#tatil", "#otelseyahat", "#keyif", "#instahotel"]
        return caption, hashtags

    # REST isteği
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    payload = {
        "contents": [{
            "parts": [
                {"text": sys_prompt},
                {"inline_data": {"mime_type": "image/jpeg", "data": b64}}
            ]
        }]
    }

    try:
        resp = requests.post(url, json=payload, timeout=GEMINI_TIMEOUT)
        resp.raise_for_status()
        data = resp.json()
        # Gemini dönüşünden metni al
        text = ""
        if "candidates" in data and data["candidates"]:
            parts = data["candidates"][0].get("content", {}).get("parts", [])
            text = "".join(p.get("text", "") for p in parts)
        # JSON yakala
        json_block = text.strip()
        # bazen kod bloğu içine koyabiliyor
        m = re.search(r"\{[\s\S]*\}", json_block)
        if m:
            json_block = m.group(0)
        # eval değil, json
        import json as pyjson
        obj = pyjson.loads(json_block)
        caption = (obj.get("caption") or "").strip()
        tags = obj.get("hashtags") or []
        # temizlik
        tags = [t.strip() for t in tags if t.strip()]
        tags = [t if t.startswith("#") else f"#{t.lstrip('#')}" for t in tags]
        # 4'e sabitle
        tags = tags[:4]
        while len(tags) < 4:
            tags.append("#tatil")
        # caption boşsa fallback
        if not caption:
            caption = "Lezzet ve keyif dolu anlar sizi bekliyor ✨🌊"
        return caption, tags
    except Exception:
        # fallback
        caption = "Doğa ve huzurun buluştuğu bu karede güzel bir gün dileriz 🌿😊"
        tags = ["#otel", "#keşfet", "#tatil", "#instatravel"]
        return caption, tags


# -----------------------------------------------------------------------------
# DOCX BUILDER
# -----------------------------------------------------------------------------
def set_page_margins(doc: Document, margin_cm: float = PAGE_MARGIN_CM):
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def para(doc: Document, text: str, size_pt: float, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.size = Pt(size_pt)
    font.bold = bold
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if align:
        p.alignment = align
    return p


def add_image_fixed_cm(doc: Document, image_bytes: bytes, width_cm: float = IMG_W_CM, height_cm: float = IMG_H_CM):
    # JPEG'e çevirip ekle (seek edilebilir buffer)
    img = open_image_safely(image_bytes)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=90)
    buf.seek(0)
    doc.add_picture(buf, width=Cm(width_cm), height=Cm(height_cm))


def shrink_caption_if_needed(text: str) -> Tuple[str, float]:
    """
    Çok uzun açıklamalarda fontu kademeli küçülterek sayfaya sığmayı kolaylaştır.
    """
    length = len(text)
    if length <= 260:
        return text, FONT_BASE_PT
    if length <= 340:
        return text, 9.5
    if length <= 420:
        return text, 9.0
    # aşırı uzunsa kısalt (sonuna …)
    t = textwrap.shorten(text, width=420, placeholder="…")
    return t, 9.0


def write_plan_docx(items: List[Dict[str, Any]], plan_name: str, hotel_info: str) -> bytes:
    """
    items: [{date: datetime.date, image_bytes: bytes, caption: str, hashtags: [str, str, str, str]}]
    """
    doc = Document()
    set_page_margins(doc, PAGE_MARGIN_CM)

    for idx, it in enumerate(items, start=1):
        # 1) Tarih (üstte, ortalı, kalın)
        date_txt = it["date"].strftime("%d.%m.%Y")
        para(doc, date_txt, FONT_DATE_PT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # 2) Görsel (16×20 cm)
        add_image_fixed_cm(doc, it["image_bytes"], IMG_W_CM, IMG_H_CM)

        # 3) Açıklama (emoji'li, başlıksız)
        cap, font_pt = shrink_caption_if_needed(it["caption"])
        para(doc, cap, font_pt)

        # 4) İletişim (ayrı paragraf, başlıksız)
        contact_clean = hotel_info.strip()
        if contact_clean:
            # tek paragrafa çok satır olacak şekilde
            para(doc, contact_clean, FONT_CONTACT_PT)

        # 5) Hashtag (tek satır, 4 adet)
        tags = it["hashtags"][:4]
        while len(tags) < 4:
            tags.append("#tatil")
        tags_line = " ".join(tags)
        para(doc, tags_line, FONT_BASE_PT)

        if idx != len(items):
            doc.add_page_break()

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


# -----------------------------------------------------------------------------
# HELPERS: DATES
# -----------------------------------------------------------------------------
def build_dates(plan_month: str, count: int, every_n_days: int) -> List[date]:
    """
    plan_month: 'YYYY-MM' (ör. '2025-09')
    """
    year, month = map(int, plan_month.split("-"))
    days_in_month = monthrange(year, month)[1]
    d = date(year, month, 1)
    dates = []
    for i in range(count):
        if d.day > days_in_month:
            # ay biterse aynı aralıkla bir sonraki aya taşı
            month += 1
            if month == 13:
                year += 1
                month = 1
            days_in_month = monthrange(year, month)[1]
            d = date(year, month, 1)
        dates.append(d)
        d = d + timedelta(days=every_n_days)
    return dates


# -----------------------------------------------------------------------------
# ROUTES: AUTH
# -----------------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    user = current_user(request)
    if user:
        return RedirectResponse("/plan")
    return RedirectResponse("/login")


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request, "app_name": APP_NAME})


@app.post("/login")
def login_action(request: Request, username: str = Form(...), password: str = Form(...)):
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT id, username, password_sha, role FROM users WHERE username=?", (username,))
    row = cur.fetchone()
    conn.close()
    if not row or row[2] != sha256(password):
        return templates.TemplateResponse(
            "login.html",
            {"request": request, "app_name": APP_NAME, "error": "Kullanıcı adı veya şifre hatalı."},
            status_code=401,
        )
    request.session["user"] = {"id": row[0], "username": row[1], "role": row[3]}
    return RedirectResponse("/plan", status_code=302)


@app.get("/logout")
def logout(request: Request):
    request.session.clear()
    return RedirectResponse("/login", status_code=302)


# -----------------------------------------------------------------------------
# ROUTES: ADMIN
# -----------------------------------------------------------------------------
@app.get("/admin", response_class=HTMLResponse)
def admin_page(request: Request, user=Depends(require_admin)):
    # mevcut kullanıcıları listele
    conn = db()
    cur = conn.cursor()
    cur.execute("SELECT id, username, role, created_at FROM users ORDER BY id DESC")
    rows = cur.fetchall()
    conn.close()
    users = [{"id": r[0], "username": r[1], "role": r[2], "created_at": r[3]} for r in rows]
    return templates.TemplateResponse(
        "admin.html",
        {"request": request, "app_name": APP_NAME, "users": users, "me": user},
    )


@app.post("/admin/users/create")
def admin_create_user(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    role: str = Form("user"),
    user=Depends(require_admin),
):
    if not username or not password:
        raise HTTPException(status_code=400, detail="Eksik alan.")
    conn = db()
    cur = conn.cursor()
    try:
        cur.execute(
            "INSERT INTO users(username, password_sha, role, created_at) VALUES(?,?,?,?)",
            (username, sha256(password), role if role in ("user", "admin") else "user", datetime.utcnow().isoformat()),
        )
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        return templates.TemplateResponse(
            "admin.html",
            {"request": request, "app_name": APP_NAME, "error": "Kullanıcı adı zaten var.", "users": []},
            status_code=400,
        )
    conn.close()
    return RedirectResponse("/admin", status_code=302)


# -----------------------------------------------------------------------------
# ROUTES: PLAN
# -----------------------------------------------------------------------------
@app.get("/plan", response_class=HTMLResponse)
def plan_page(request: Request, user=Depends(require_login)):
    return templates.TemplateResponse("plan.html", {"request": request, "app_name": APP_NAME, "me": user})


@app.post("/api/plan")
async def api_plan(
    request: Request,
    images: List[UploadFile] = File(...),
    plan_month: str = Form(...),          # 'YYYY-MM'
    every_n_days: int = Form(2),
    plan_name: str = Form("Instagram_Plani"),
    hotel_info: str = Form(""),
    user=Depends(require_login),
):
    # görselleri oku
    raw_images: List[bytes] = []
    for f in images:
        raw = await f.read()
        raw_images.append(raw)

    # tarihleri üret
    dates = build_dates(plan_month, len(raw_images), every_n_days)

    # her görsel için caption & hashtag (4)
    items = []
    for raw, d in zip(raw_images, dates):
        caption, tags = gemini_json(raw, hotel_info)
        # 4'e sabitle
        tags = tags[:4]
        while len(tags) < 4:
            tags.append("#tatil")
        items.append({
            "date": d,
            "image_bytes": raw,
            "caption": caption,
            "hashtags": tags,
        })

    # docx yaz
    bin_docx = write_plan_docx(items, plan_name, hotel_info)

    # çıktı adı
    today_str = datetime.now().strftime("%Y%m%d-%H%M")
    safe_name = re.sub(r"[^A-Za-z0-9_\-şŞıİçÇöÖüÜğĞ ]+", "", plan_name).strip().replace(" ", "_")
    filename = f"{safe_name}_{today_str}.docx"
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"'
    }
    return StreamingResponse(io.BytesIO(bin_docx), headers=headers, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# -----------------------------------------------------------------------------
# HEALTH
# -----------------------------------------------------------------------------
@app.get("/health")
def health():
    return {"ok": True, "db": USAGE_DB_PATH, "app": APP_NAME}
